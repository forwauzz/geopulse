from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "tmp" / "pdfs" / "linkedin-ai-visibility-playbook"
OUTPUT = (
    ROOT
    / "output"
    / "documents"
    / "geopulse-ai-visibility-reporting-playbook-linkedin.docx"
)

PAGE_WIDTH = Inches(7.5)
PAGE_HEIGHT = Inches(9.375)
# A tiny vertical allowance prevents Word or LibreOffice from pushing a
# full-bleed inline image onto the following page because of its text baseline.
IMAGE_WIDTH = Inches(7.5)
IMAGE_HEIGHT = Inches(9.25)


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(1)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    doc.core_properties.title = "The AI Visibility Reporting Playbook for Agencies"
    doc.core_properties.subject = (
        "A practical workflow for measured AI visibility and readiness"
    )
    doc.core_properties.author = "GEO-Pulse"
    doc.core_properties.keywords = "AI visibility, GEO, agency reporting"

    page_descriptions = [
        "Cover: Stop reporting AI visibility by hand.",
        "A repeatable measure, diagnose, fix, and verify workflow.",
        "Start with the exact buyer questions that matter.",
        "Show the denominator and preserve measurement scope.",
        "The GEO-Pulse 24-check readiness framework.",
        "Five technical access gates to verify first.",
        "Seven questions for a client-ready report.",
        "Closing page: Measure. Fix. Verify.",
    ]

    for page_number in range(1, 9):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run()
        image_path = SOURCE_DIR / f"page-{page_number}.png"
        shape = run.add_picture(
            str(image_path), width=IMAGE_WIDTH, height=IMAGE_HEIGHT
        )
        set_alt_text(
            shape,
            f"GEO-Pulse playbook page {page_number}",
            page_descriptions[page_number - 1],
        )

        if page_number < 8:
            run.add_break(WD_BREAK.PAGE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
